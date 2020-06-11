# -*- coding=utf-8 -*-
# any issue , please contact yuan.fei@mediatek.com

import re
import os
import sys


from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor


RE_COV = re.compile(r"\s*COV=(\w+)\s+NAME=((\w|\/)+)\:\:(\w+|[|])")
RE_INST = re.compile(r"((\w|\/)+)\/(\w+)")
#RE_DOC = re.compile(r"(\w+)\.(\w+)")
cov_list_header = ("name","coverage","total","covered","missing","weight")

TYPE_OP = { 'TYPE': 0, 'INST':1, 'ALL_INST':2 }

class cov_doc_proc:
	#doc = None
	def __init__(self, doc_name, debug):
		self.file_name=doc_name
		self.doc = Document(self.file_name)
		self.debug = debug

	def tabBgColor(self, table,cols,colorStr):
	    shading_list = locals()
	    for i in range(cols):
	        shading_list['shading_elm_'+str(i)] = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor = colorStr))
	        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])
	
	def move_table_after(self, table, paragraph):
		tbl, p = table._tbl, paragraph._p
		p.addnext(tbl)
	
	
	def proc_doc(self,  cov_dict, prefix):
		
		#styles = doc.styles
		#for s in styles:
		#	if s.type == WD_STYLE_TYPE.TABLE:
		#		print("table style ", s.name)
	
		for line in self.doc.paragraphs:
			#print (line.text)
			cov_re = RE_COV.search(line.text)
			if(cov_re != None):
				print(f"Find COV : {cov_re.group()}")
				c_type = cov_re.group(1)
				c_path_org = cov_re.group(2)
				c_name = cov_re.group(4)
				row_num =2
	
				if c_type in TYPE_OP:
					cov_mode = TYPE_OP[c_type]
				else:
					print(f"FATAL :: illegal COV_MODE {c_type}")
					exit(1)
	
				if cov_mode==1 :
					inst_re = RE_INST.search(c_path_org)
					c_path = inst_re.group(1)
					c_inst       = inst_re.group(3)
					full_name = c_path_org+'::'+c_name
				else :
					c_path = c_path_org
					c_inst = 'NA'
					full_name = c_path+'::'+c_name
	
				if c_path in cov_dict :
					cg_dict = cov_dict[c_path]
					cg_type_dict = cg_dict['TYPE']
					cg_insts_dict = cg_dict['INST']
				else:
					print ("Error : covergroup %s not found in xml" % c_path)
					continue
					#exit(1)
					
				if cov_mode == 0:
					tmp_dict = cg_type_dict
				elif cov_mode ==1:
					tmp_dict = cg_insts_dict[c_inst]
				elif cov_mode ==2:
					tmp_dict = cg_type_dict
					row_num = row_num+len(cg_insts_dict) 
	
				if c_name in tmp_dict :
					item_dict = tmp_dict[c_name]
				else:
					print ("Error : cover_name %s not found in xml" % c_name)
					font = line.add_run("\nError : coverage \"%s\" not found in xml" % c_name).font
					font.color.rgb = RGBColor(255, 0, 0)
					continue
					#exit(1)
	
				if self.debug:
					print(f"type : {c_type}")
					print(f"path : {c_path}")
					print(f"inst : {c_inst}")
					print(f"NAME : {c_name}")
				table=self.doc.add_table(row_num,6)
				table.autofit=True
				table.style='Table Grid'
			
				#  GREEN #32CD32   100
				#  ORINGE #EE9A00  50~100
				#  RED	#FF0000    <50
				#  blue #00BFFF
				#  GRAY #696969
				  
				weight =  int( item_dict['weight'] )
				#print ("weight=",weight)
				if weight == 0:
					colorStr = 'C0C0C0'
				else :
					score = float( item_dict['coverage'] )
					if score == 100:
						colorStr = '32CD32'
					elif score > 50:
						colorStr = 'EE9A00'
					else :
						colorStr = 'FF0000'
	
				self.tabBgColor(table, 6, colorStr )
				for i in range(6):
					## header
					cell = table.cell(0,i)
					if i!=0:
						cell.text = cov_list_header[i]
					else :
						if prefix is None :
							cell.text = cov_list_header[i]
						else :
							cell.text = cov_list_header[i] + " ["+prefix+"]"
					## sccond 
					cell_fill = table.cell(1,i)
	
					if i!=0:
						cell_fill.text = item_dict[cov_list_header[i]]
					else :
						cell_fill.text = full_name
					if cov_mode==2:
						keys = list(cg_insts_dict.keys())
						cnt=0
						for key in keys:
							tmp_inst_dict = cg_insts_dict[key]
							
							tmp_inst_item_dict = tmp_inst_dict[c_name]
						
							cell_inst = table.cell(2+cnt,i)
							if i!=0:
								cell_inst.text = tmp_inst_item_dict[cov_list_header[i]]
							else :
								cell_inst.text = c_path+'/'+key+'::'+c_name
							cnt=cnt+1
	
				self.move_table_after(table, line)
	
		out_file = '[ANNOTATE]'+self.file_name
		self.doc.save(out_file)

